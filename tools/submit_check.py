"""投稿する成果物（docx / PDF）が投稿システムを通る形かを確かめる。

なぜ要るか
  2026-09-01 の bioRxiv 投稿で、docx から PDF への変換が**サーバ側で失敗した**。
  原因は画像幅 6.5 inch が本文幅 6.0 inch（用紙 8.5 − 左右余白 1.25 × 2）を
  はみ出していたことで、原稿を読んでも気づけない。エラーメッセージも
  「conversion failed」だけで、どこが悪いかは教えてくれなかった。

  投稿の最終工程で初めて分かる不備は、そこまでの全工程をやり直させる。
  ここで先に落とす。

見るもの
  投稿先が著者を受け付けるか  所属要件・endorsement・既出可否（今回落ちたのはここ）
  用紙サイズと向き        bioRxiv は 8.5 x 11 inch の portrait を要求する
  画像が本文幅に収まるか   はみ出すと変換系が落ちる
  画像・表が入っているか   図表が抜けた docx を出すと投稿が差し戻される
  PDF が暗号化されていないか  locked PDF は受理されない
  docx/PDF が md より新しいか  作り直し忘れは内容検査では見つからない
  PDF の必須要素          Title / Author / Abstract が 1 ページ目にあるか

使い方
  pixi run submit-check
"""
from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPTS = ROOT / "manuscript"

# bioRxiv の要求（docs/計画と運用/18 の実施記録）
PAGE_W, PAGE_H = 8.5, 11.0
TOL = 0.05

# (表示名, md の原本, docx, pdf)
TARGETS = [
    ("英文（投稿版）",
     MANUSCRIPTS / "04-preprint-v1-en.md",
     MANUSCRIPTS / "04-preprint-v1-en.docx",
     MANUSCRIPTS / "04-preprint-v1-en.pdf"),
    ("和文",
     MANUSCRIPTS / "02-投稿原稿_日本語_v1.md",
     MANUSCRIPTS / "02-投稿原稿_日本語_v1.docx",
     None),
]


# 投稿先ごとの受理条件。**一次情報で確認した日付を必ず添える。**
# 2026-09-01 の bioRxiv 拒否は、ここを確認せずに投稿したために起きた。
# 「原稿と投稿物の形」が整っていても、投稿先が著者を受け付けないなら落ちる。
VENUE = "osf"   # 投稿先を変えたらここを変える

VENUE_RULES = {
    "biorxiv": {
        "label": "bioRxiv",
        "requires_org_affiliation": True,
        "allows_prior_preprint": False,
        "verified": "2026-09-02",
        "why": ("組織による oversight を要求する（研究不正の申し立て先が必要）。"
                "2026-09-01 に Affiliation = Independent Researcher で拒否された。"
                "他のプレプリントサーバに既出の原稿も受け付けない"),
    },
    "osf": {
        "label": "OSF Preprints",
        "requires_org_affiliation": False,
        "allows_prior_preprint": True,
        "verified": "2026-09-02",
        "why": "所属は任意（OSF Support で確認）。DOI と永続 URL が付く。PDF 推奨",
    },
    "arxiv": {
        "label": "arXiv",
        "requires_org_affiliation": False,
        "allows_prior_preprint": True,
        "needs_endorsement": True,
        "verified": "2026-09-02",
        "why": ("所属は不要だが、分野への初投稿には既存 arXiv 著者からの endorsement が要る。"
                "2026-01-21 に方針が厳格化し、所属メールは資格として認められなくなった"),
    },
}

# 組織による oversight を提供しない所属の書き方。ここに該当すると bioRxiv 系は落ちる。
NO_ORG_AFFILIATION = ("independent researcher", "independent scholar",
                      "unaffiliated", "no affiliation", "private researcher")


def check_eligibility(md: Path, venue: str) -> tuple[list[str], list[str]]:
    """投稿先が著者を受け付けるか。原稿の形ではなく投稿資格を見る。"""
    errors, notes = [], []
    rule = VENUE_RULES.get(venue)
    if rule is None:
        return [f"投稿先 {venue!r} の受理条件が未登録。一次情報で確認して VENUE_RULES に足す"], []
    notes.append(f"投稿先 {rule['label']}（受理条件の確認日 {rule['verified']}）")
    if not md.exists():
        return errors, notes

    text = md.read_text(encoding="utf-8-sig")
    m = re.search(r"^\*\*Affiliation\*\*\s*(.+)$", text, re.M)
    if m is None:
        m = re.search(r"^\*\*所属\*\*\s*(.+)$", text, re.M)
    if m is None:
        errors.append("原稿に Affiliation 行がない")
        return errors, notes
    aff = m.group(1).strip()
    notes.append(f"Affiliation: {aff}")

    if rule["requires_org_affiliation"] and any(k in aff.lower() for k in NO_ORG_AFFILIATION):
        errors.append(f"{rule['label']} は組織の所属を要求するが、Affiliation が "
                      f"「{aff}」になっている → {rule['why']}")
    if rule.get("needs_endorsement"):
        notes.append(f"注意: {rule['label']} は endorsement が要る。投稿前に取得しておく")
    if not rule["allows_prior_preprint"]:
        notes.append(f"注意: {rule['label']} は他サーバに既出の原稿を受け付けない。"
                     "先に別サーバへ出していないか確認する")
    return errors, notes


def check_freshness(md: Path, docx: Path, pdf: Path | None) -> list[str]:
    """投稿物が原稿より古くないか。

    md を直して docx を作り直すのを忘れると、修正前の原稿を投稿することになる。
    これはどんな内容検査でも見つからない（docx 自体は整合している）。
    """
    errors = []
    if not md.exists():
        return [f"md の原本がない: {md.name}"]
    for label, out in (("docx", docx), ("PDF", pdf)):
        if out is None or not out.exists():
            continue
        if out.stat().st_mtime < md.stat().st_mtime - 1:
            errors.append(f"{label} が md より古い（{out.name}）"
                          " → pixi run docx で作り直してから投稿する")
    return errors


def check_docx(path: Path) -> tuple[list[str], list[str]]:
    errors, notes = [], []
    try:
        from docx import Document
    except ImportError:
        return ["python-docx が無いので docx を検査できない"], []
    if not path.exists():
        return [f"docx が無い: {path.name}"], []

    d = Document(str(path))
    s = d.sections[0]
    w, h = s.page_width.inches, s.page_height.inches
    body_w = w - s.left_margin.inches - s.right_margin.inches
    notes.append(f"用紙 {w:.2f} x {h:.2f} inch / 本文幅 {body_w:.2f} inch")

    if abs(w - PAGE_W) > TOL or abs(h - PAGE_H) > TOL:
        errors.append(f"用紙が {w:.2f} x {h:.2f} inch。bioRxiv は {PAGE_W} x {PAGE_H} を要求する")
    if h <= w:
        errors.append("landscape になっている。portrait が要る")

    over = [(i, sh.width.inches) for i, sh in enumerate(d.inline_shapes, 1)
            if sh.width.inches > body_w + 0.01]
    if over:
        for i, iw in over:
            errors.append(f"画像 {i} の幅 {iw:.2f} inch が本文幅 {body_w:.2f} inch を超えている"
                          " → PDF 変換が落ちる")
    notes.append(f"埋め込み画像 {len(d.inline_shapes)} 点、表 {len(d.tables)} 点")
    if not d.inline_shapes:
        errors.append("画像が 1 点も埋め込まれていない")
    if not d.tables:
        errors.append("表が 1 点も入っていない")

    # 埋め込みメディアの形式（bmp / psd / xls は投稿系が受け付けない）
    bad_ext = {".bmp", ".psd", ".xls", ".cdr", ".pict"}
    with zipfile.ZipFile(path) as z:
        media = [n for n in z.namelist() if "/media/" in n]
        for n in media:
            if Path(n).suffix.lower() in bad_ext:
                errors.append(f"受け付けられない画像形式が埋まっている: {Path(n).name}")
    notes.append(f"docx サイズ {path.stat().st_size / 1024 / 1024:.2f} MB")
    return errors, notes


def check_pdf(path: Path) -> tuple[list[str], list[str]]:
    errors, notes = [], []
    if path is None:
        return [], []
    if not path.exists():
        return [f"PDF が無い: {path.name}（投稿には PDF を出す方が確実）"], []
    try:
        from pypdf import PdfReader
    except ImportError:
        return ["pypdf が無いので PDF を検査できない"], []

    r = PdfReader(str(path))
    notes.append(f"PDF {len(r.pages)} ページ、{path.stat().st_size / 1024 / 1024:.2f} MB")
    if r.is_encrypted:
        errors.append("PDF が暗号化されている。locked PDF は受理されない")
    box = r.pages[0].mediabox
    w, h = float(box.width) / 72, float(box.height) / 72
    notes.append(f"PDF 用紙 {w:.2f} x {h:.2f} inch")
    if abs(w - PAGE_W) > TOL or abs(h - PAGE_H) > TOL:
        errors.append(f"PDF の用紙が {w:.2f} x {h:.2f} inch")

    imgs = sum(len(p.images) for p in r.pages)
    notes.append(f"PDF 内の画像 {imgs} 点")
    if imgs == 0:
        errors.append("PDF に画像が 1 点も無い（図が抜けている）")

    full = "".join((p.extract_text() or "") for p in r.pages)
    first = r.pages[0].extract_text() or ""
    for key, where, blob in (("Abstract", "1 ページ目", first),
                             ("Ryota Ueda", "1 ページ目", first),
                             ("orcid", "本文", full.lower()),
                             ("Competing interests", "本文", full),
                             ("Data and code availability", "本文", full)):
        if key.lower() not in blob.lower():
            errors.append(f"PDF の{where}に「{key}」が見つからない")

    if "will be given" in full:
        errors.append("PDF に「DOI will be given」が残っている。DOI を記入してから出す")
    return errors, notes


def main() -> int:
    total = 0
    for label, md, docx, pdf in TARGETS:
        print(f"=== 投稿物の検査: {label} ===")
        ev, nv = check_eligibility(md, VENUE)
        e0 = check_freshness(md, docx, pdf)
        e1, n1 = check_docx(docx)
        e2, n2 = check_pdf(pdf)
        for n in nv + n1 + n2:
            print(f"  - {n}")
        errs = ev + e0 + e1 + e2
        if errs:
            total += len(errs)
            print(f"\n  【要修正 {len(errs)} 件】")
            for e in errs:
                print(f"    x {e}")
        else:
            print("  投稿物の問題は見つからなかった")
        print()
    if total:
        print(f"要修正 {total} 件")
        return 1
    print("投稿できる形になっている")
    return 0


if __name__ == "__main__":
    sys.exit(main())
