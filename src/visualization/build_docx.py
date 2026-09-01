"""原稿の Markdown を Word (.docx) に変換する（査読・コメント用）。

原本は manuscript/*.md 側に置き、docx はそこから生成する使い捨てにする。
これにより「Word で直した内容が原本に反映されない」という事故を避ける。

対応する記法は原稿で実際に使うものだけに絞った。
  # / ## / ###  見出し
  段落
  - 箇条書き / 1. 番号つき
  | 表 |（区切り行 |---| を含む）
  **強調**
  ---（水平線 = セクション区切りとして無視）

使い方:
  pixi run docx                       manuscript/*.md をすべて変換
  pixi run python -m src.visualization.build_docx path/to/file.md
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from ..common import ROOT

JP_FONT = "Yu Gothic"
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
IMG_RE = re.compile(r"^!\[(?P<cap>.*?)\]\((?P<path>[^)]+)\)\s*$")
# Letter・余白 1 インチの本文幅。figures.py の DOC_WIDTH_IN と一致させる。
# 図の保存幅がこれより大きいと、その分だけ縮小されて図中の文字が読めなくなる。
IMG_WIDTH_IN = 6.0   # 本文幅（8.5 - 左右余白 1.25 x 2）に収める。はみ出すと変換系が落ちる


def _set_font(run, size: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = JP_FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # 日本語フォントは eastAsia 属性にも設定しないと反映されない
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", JP_FONT
    )


def _add_rich(paragraph, text: str, size: float = 10.5) -> None:
    """**強調** を太字 run に分けて追加する。"""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            _set_font(paragraph.add_run(text[pos:m.start()]), size)
        _set_font(paragraph.add_run(m.group(1)), size, bold=True)
        pos = m.end()
    if pos < len(text):
        _set_font(paragraph.add_run(text[pos:]), size)


def _split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def convert(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8-sig").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = JP_FONT
    style.font.size = Pt(10.5)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line or line.strip() == "---":
            i += 1
            continue

        # 図: ![キャプション](相対パス)
        m = IMG_RE.match(line)
        if m:
            img = (md_path.parent / m.group("path")).resolve()
            if img.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(img), width=Inches(IMG_WIDTH_IN))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cap.paragraph_format.space_after = Pt(12)
                _add_rich(cap, m.group("cap"), size=9)
                for r in cap.runs:
                    r.font.color.rgb = RGBColor.from_string("525252")
            else:
                p = doc.add_paragraph()
                _add_rich(p, f"［図が見つかりません: {m.group('path')}］")
            i += 1
            continue

        # 表: ヘッダ行 + 区切り行 + データ行
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:\-|]+\|$", lines[i + 1].strip()):
            header = _split_row(line)
            rows = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append(_split_row(lines[j]))
                j += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for k, h in enumerate(header):
                cell = table.rows[0].cells[k]
                cell.text = ""
                _add_rich(cell.paragraphs[0], h, size=9.5)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for row in rows:
                cells = table.add_row().cells
                for k, v in enumerate(row[:len(header)]):
                    cells[k].text = ""
                    _add_rich(cells[k].paragraphs[0], v, size=9.5)
            doc.add_paragraph()
            i = j
            continue

        # 見出し
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level, text = len(m.group(1)), m.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
            p.paragraph_format.space_after = Pt(4)
            size = {1: 18, 2: 14, 3: 12, 4: 11}[level]
            _add_rich(p, text, size=size)
            for r in p.runs:
                r.bold = True
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # 箇条書き
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_rich(p, m.group(2))
            i += 1
            continue
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_rich(p, m.group(2))
            i += 1
            continue

        # 段落（連続行は 1 段落にまとめる）
        buf = [line]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not re.match(r"^(#{1,4}\s|[-*]\s|\d+\.\s|\|)", lines[j]) and lines[j].strip() != "---":
            buf.append(lines[j].strip())
            j += 1
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _add_rich(p, "".join(buf))
        i = j

    doc.save(out_path)
    print(f"  {out_path.name}")


def main(argv: list[str] | None = None) -> int:
    args = argv if argv is not None else sys.argv[1:]
    targets = [Path(a) for a in args] if args else sorted((ROOT / "manuscript").glob("*.md"))
    if not targets:
        print("変換対象の Markdown がない")
        return 1
    # 書き込めるかを先に全部確かめる。
    # Word で開いたままの .docx は上書きできず、途中まで生成して止まると
    # 「一部だけ新しい docx」ができる。実際にそれをレビュアーに送ってしまい、
    # 解決済みの指摘が再び返ってきた。1 つでもロックされていたら何も書かない。
    locked = []
    for md in targets:
        out = md.with_suffix(".docx")
        if not out.exists():
            continue
        try:
            with out.open("r+b"):
                pass
        except OSError:
            locked.append(out.name)
    if locked:
        print("★ 次のファイルが開かれていて上書きできない。閉じてから再実行する:")
        for name in locked:
            print(f"    {name}")
        print("  1 つでもロックされていると一部だけ古い docx が残るので、何も書かない。")
        return 1

    for md in targets:
        convert(md, md.with_suffix(".docx"))
    return 0


if __name__ == "__main__":
    sys.exit(main())
